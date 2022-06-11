import functools
from os import listdir
from os.path import isfile, join

from docx import Document
from xml.dom.minidom import parse, parseString

WORD_NAMESPACE = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
TEXT = WORD_NAMESPACE + "t"


def _get_accepted_text(p, debug=False):
    """Return text of a paragraph after accepting all changes"""
    xml = p._p.xml
    parsed_texts = []
    count = 0

    for node in parseString(xml.replace('\n', '')).firstChild.childNodes:

        if debug:
            count += 1
            print("*** NODE {0} ***".format(count))

        parsed_texts.append(_recursive_node_text(node, level=0 if debug else -1))

    return parsed_texts


def _recursive_node_text(node, input_text="", accepted_text="", level=0):

    if level != -1:
        level += 1
        print("{0}> {1}".format("".join(['-' for x in range(0, level)]), node.nodeName))

    if node.nodeName in ['w:pPr', 'w:rPr', 'w:bookmarkStart', 'w:bookmarkEnd', 'w:spacing', 'w:lastRenderedPageBreak',
                         'w:br', 'w:fldChar', 'w:instrText', 'w:tab', 'w:noBreakHyphen', 'w:proofErr', 'w:pict',
                         'w:commentRangeStart',  'w:commentRangeEnd', 'w:commentReference']:
        r1, r2 = input_text, accepted_text
    elif len(node.childNodes) > 0:
        for childNode in node.childNodes:
            it, at = _recursive_node_text(childNode, input_text, accepted_text, level)
            input_text = it
            accepted_text = at

        r1, r2 = input_text, accepted_text
    else:
        if node.parentNode.nodeName == 'w:t' and node.parentNode.parentNode.parentNode.nodeName == 'w:ins':
            r1, r2 = input_text, accepted_text + node.wholeText
        elif node.parentNode.nodeName == 'w:delText':
            r1, r2 = input_text + node.wholeText, accepted_text
        elif node.parentNode.nodeName == 'w:t':
            r1, r2 = input_text + node.wholeText, accepted_text + node.wholeText
        else:
            if node.wholeText == True:
                to_append = node.wholeText
            else:
                to_append = ""

            r1 = input_text + to_append
            r2 = accepted_text + to_append

    if level != -1:
        level -= 1
        print("<-{0} RETURN FROM {1}".format("".join(['-' for x in range(0, level)]), node.nodeName))

    return r1, r2


class DocumentParser(object):

    def __init__(self, documents_path):
        self.DOCS_PATH = documents_path

    def process_files(self, verbose=False):

        sentence_pairs = []

        sorted_docs = listdir(self.DOCS_PATH)
        sorted_docs.sort()

        for document_name in [directory_file for directory_file in sorted_docs
                              if isfile(join(self.DOCS_PATH, directory_file))]:
            print("Processing File: " + self.DOCS_PATH + document_name)
            document = Document(self.DOCS_PATH + document_name)

            count = 0
            for p in document.paragraphs:
                count += 1
                if (verbose):
                  print("Processing Paragraph {0}/{1}".format(count, len((document.paragraphs))))

                parsed_texts = _get_accepted_text(p)
                source_text, validated_text = functools.reduce(lambda x, y: (x[0] + y[0], x[1] + y[1]), parsed_texts) \
                    if len(parsed_texts) > 0 else ("", "")

                if (verbose):
                    print(source_text)
                    print(validated_text)

                sentence_pairs.append((source_text, validated_text))

        return sentence_pairs